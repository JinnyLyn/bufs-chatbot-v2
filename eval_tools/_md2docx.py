"""Minimal Markdown -> .docx for the two reports (headings, tables, bold, bullets,
blockquotes, hr). Korean font. Not a general MD parser — just what the reports use."""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTS_DIR = "reports"
DEFAULT_FILES = [f"{REPORTS_DIR}/REPORT_결과.md", f"{REPORTS_DIR}/REPORT_vs_BUFS.md",
                 f"{REPORTS_DIR}/REPORT_장애대응.md"]


def add_runs(p, text):
    # split on **bold** and `code`
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"
        else:
            p.add_run(tok)


def set_font(doc, name="Malgun Gothic"):
    st = doc.styles["Normal"]
    st.font.name = name; st.font.size = Pt(10.5)
    # east-asian font
    from docx.oxml.ns import qn
    st.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def convert(md_path):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document(); set_font(doc)
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1; continue
        # table block
        if ln.lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i].strip()); i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in block
                    if not re.fullmatch(r"[\|\-: ]+", r)]
            if rows:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=ncol); t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    cells = t.add_row().cells
                    for ci in range(ncol):
                        cells[ci].text = ""
                        p = cells[ci].paragraphs[0]
                        add_runs(p, row[ci] if ci < len(row) else "")
                        if ri == 0:
                            for rr in p.runs: rr.bold = True
            continue
        # headings
        m = re.match(r"(#{1,4})\s+(.*)", ln)
        if m:
            lvl = len(m.group(1)); doc.add_heading(m.group(2).strip(), level=min(lvl, 4)); i += 1; continue
        # hr
        if re.fullmatch(r"-{3,}", ln.strip()):
            doc.add_paragraph().add_run("─" * 30); i += 1; continue
        # blockquote
        if ln.lstrip().startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            r = p.add_run("▏ "); r.bold = True
            add_runs(p, ln.lstrip()[1:].strip().lstrip(">").strip()); i += 1; continue
        # bullet
        if re.match(r"\s*[-*]\s+", ln):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, re.sub(r"^\s*[-*]\s+", "", ln)); i += 1; continue
        # numbered
        if re.match(r"\s*\d+\.\s+", ln):
            p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\s*\d+\.\s+", "", ln)); i += 1; continue
        # paragraph
        p = doc.add_paragraph(); add_runs(p, ln); i += 1

    out = md_path.rsplit(".", 1)[0] + ".docx"
    doc.save(out); return out


if __name__ == "__main__":
    # Usage: python eval_tools/_md2docx.py [report.md ...]  (default: every report in reports/)
    for f in sys.argv[1:] or DEFAULT_FILES:
        print("wrote", convert(f))
